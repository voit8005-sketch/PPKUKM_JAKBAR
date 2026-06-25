# -*- coding: utf-8 -*-
"""
Script untuk menghasilkan Bab 3 Laporan Kerja Praktek
Website Sudin PPKUKM Jakarta Barat
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(4.0)
section.right_margin  = Cm(3.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(3.0)

# ─────────────────────────────────────────────
# STYLE HELPERS
# ─────────────────────────────────────────────
def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    run = p.add_run(text)
    sizes = {1: 14, 2: 13, 3: 12}
    set_font(run, size=sizes.get(level, 12), bold=True)
    return p

def body(text, indent=0, justify=True, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def bullet(text, indent=1.0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def numbered(text, indent=1.0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(10)
    run = p.add_run(text)
    set_font(run, size=11, italic=True)
    return p

def add_table_header_row(table, headers, col_widths=None):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        set_font(run, size=11, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)
    if col_widths:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Cm(w)

def add_data_row(table, row_idx, data, center_cols=None):
    row = table.rows[row_idx]
    center_cols = center_cols or []
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(str(val))
        set_font(run, size=11)
        # Zebra
        if row_idx % 2 == 0:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EBF5FB')
            tcPr.append(shd)

def diagram_box(lines, title=None):
    """Render sebuah diagram ASCII di dalam kotak."""
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after  = Pt(2)
        r = p.add_run(f'[ {title} ]')
        set_font(r, size=11, bold=True)

    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        r = p.add_run(line)
        set_font(r, name='Courier New', size=9)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  BAB 3 — PEMBAHASAN
# ══════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(18)
r = p_title.add_run('BAB III\nPEMBAHASAN')
set_font(r, size=14, bold=True)

# ──────────────────────────────────────────────────────────
# 3.1 Gambaran Umum Sistem
# ──────────────────────────────────────────────────────────
heading('3.1 Gambaran Umum Sistem', level=1)
body(
    'Website Suku Dinas Perindustrian, Perdagangan, Koperasi, Usaha Kecil dan Menengah '
    '(Sudin PPKUKM) Kota Administrasi Jakarta Barat merupakan sistem informasi berbasis web '
    'yang dirancang untuk menyediakan layanan informasi publik sekaligus mengelola konten '
    'berita/kegiatan secara internal. Sistem ini dibangun menggunakan framework Flask (Python) '
    'di sisi back-end, Tailwind CSS di sisi front-end, serta SQLite sebagai basis data '
    'default dengan dukungan MySQL untuk lingkungan produksi.'
)
body(
    'Sistem memiliki dua lapisan pengguna, yaitu pengguna publik (pengunjung umum) yang '
    'hanya dapat mengakses informasi yang telah dipublikasikan, serta pengguna internal '
    '(Admin dan Staff) yang memiliki hak akses untuk mengelola konten. Arsitektur ini '
    'memastikan bahwa informasi yang tampil di halaman publik telah melalui proses verifikasi '
    'dan persetujuan oleh administrator.'
)

# ──────────────────────────────────────────────────────────
# 3.2 Fungsi Sistem
# ──────────────────────────────────────────────────────────
heading('3.2 Fungsi Sistem', level=1)
body(
    'Sistem website Sudin PPKUKM Jakarta Barat memiliki sejumlah fungsi utama yang '
    'dibagi berdasarkan level akses pengguna. Berikut adalah penjelasan lengkap setiap '
    'fungsi yang tersedia dalam sistem:'
)

heading('3.2.1 Fungsi untuk Pengunjung Umum (Publik)', level=2)
body('Pengguna publik dapat mengakses fungsi-fungsi berikut tanpa perlu login:')
bullet('Melihat Halaman Beranda — Menampilkan berita terbaru, ticker berita berjalan, statistik UKM, dan filter kategori berita (Semua, UKM, Pelatihan, Promosi, Event).')
bullet('Melihat Profil Organisasi — Menampilkan struktur organisasi Sudin PPKUKM beserta informasi kepala dinas dan sub-seksi yang ada.')
bullet('Melihat Sub-Seksi Organisasi — Pengguna dapat mengakses halaman detail masing-masing seksi: Subbagian Tata Usaha, Seksi Perindustrian, Seksi Perdagangan, dan Seksi Koperasi & UKM.')
bullet('Melihat Detail Berita — Menampilkan isi lengkap berita beserta galeri foto dan video yang diunggah.')
bullet('Melihat Layanan — Menampilkan informasi layanan termasuk program Pelatihan Jakpreneur.')
bullet('Melihat Manfaat Pelatihan — Menampilkan manfaat program pelatihan Jakpreneur secara rinci.')
bullet('Melihat Berita Khusus (Statis) — Menampilkan dua berita unggulan statis: Bimtek Sertifikat Halal dan Semarak Bazar Pangan Ramadan.')

heading('3.2.2 Fungsi untuk Staff', level=2)
body('Staff merupakan pegawai internal yang memiliki akun khusus untuk mengelola konten. Fungsi yang tersedia:')
bullet('Login — Staff masuk ke sistem menggunakan email dan password yang diberikan oleh admin.')
bullet('Ganti Password — Staff diwajibkan mengganti password sementara pada login pertama dan dapat mengganti password kapan saja.')
bullet('Dashboard Staff — Menampilkan ringkasan berita yang dibuat oleh staff beserta status masing-masing (Pending, Dipublikasi, Ditolak).')
bullet('Buat Berita Baru — Staff dapat membuat artikel berita lengkap dengan judul, isi, kategori, tanggal kegiatan, jumlah peserta, lokasi, serta mengunggah foto (maks. 10 MB/foto) dan video (maks. 200 MB/video).')
bullet('Edit Berita (Pending/Ditolak) — Staff dapat mengedit berita yang masih berstatus pending atau ditolak sebelum diajukan ulang.')
bullet('Hapus Berita Sendiri — Staff dapat menghapus berita yang ia buat sendiri, beserta seluruh file media terkait.')
bullet('Logout — Staff dapat keluar dari sesi aktif.')

heading('3.2.3 Fungsi untuk Admin', level=2)
body('Admin merupakan pengguna dengan hak akses tertinggi. Seluruh fungsi staff dimiliki admin, ditambah dengan fungsi-fungsi berikut:')
bullet('Dashboard Admin — Menampilkan seluruh berita yang telah dipublikasikan serta jumlah berita yang menunggu review.')
bullet('Tambah Berita Langsung — Admin dapat membuat berita baru yang langsung berstatus "published" tanpa perlu melalui antrian review.')
bullet('Edit Berita (Semua Status) — Admin dapat mengedit berita milik staff maupun berita yang sudah dipublikasikan.')
bullet('Hapus Berita — Admin dapat menghapus berita apapun beserta seluruh media terkait.')
bullet('Review Berita — Admin meninjau daftar berita berstatus "pending" dari staff, kemudian memilih untuk menyetujui (publish) atau menolak disertai alasan.')
bullet('Kelola Staff — Admin dapat menambah akun staff baru, menonaktifkan/mengaktifkan kembali akun staff, serta mereset password staff.')
bullet('Pratinjau Berita Non-Publik — Admin dapat melihat berita yang belum/tidak dipublikasikan.')

# ──────────────────────────────────────────────────────────
# 3.3 Menu Utama Sistem
# ──────────────────────────────────────────────────────────
heading('3.3 Menu Utama Sistem', level=1)
body(
    'Sistem memiliki dua jenis struktur navigasi: navigasi publik yang dapat diakses oleh '
    'seluruh pengunjung, serta navigasi internal yang hanya muncul setelah pengguna berhasil '
    'login. Berikut adalah rincian menu yang tersedia:'
)

heading('3.3.1 Menu Navigasi Publik', level=2)
body('Menu ini terdapat pada bagian header (navigation bar) dan footer yang dapat diakses oleh semua pengunjung:')

tbl1 = doc.add_table(rows=6, cols=3)
tbl1.style = 'Table Grid'
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl1, ['No', 'Menu', 'Keterangan'], col_widths=[1.0, 3.5, 9.5])
data1 = [
    ('1', 'Beranda', 'Halaman utama yang menampilkan berita terkini, ticker, statistik UKM, dan filter kategori berita.'),
    ('2', 'Profil Organisasi', 'Menampilkan struktur dan informasi organisasi Sudin PPKUKM Jakarta Barat.'),
    ('3', 'Layanan', 'Informasi layanan publik termasuk program Pelatihan Jakpreneur.'),
    ('4', 'Manfaat Pelatihan', 'Rincian manfaat mengikuti program pelatihan kewirausahaan Jakpreneur.'),
    ('5', 'Admin Login', 'Tombol masuk khusus staf/admin yang mengarah ke halaman login (hanya tampil jika belum login).'),
]
for i, row_data in enumerate(data1, 1):
    add_data_row(tbl1, i, row_data, center_cols=[0])
caption('Tabel 3.1 Menu Navigasi Publik')

heading('3.3.2 Menu Navigasi Profil Organisasi', level=2)
body('Di dalam halaman Profil Organisasi, terdapat sub-navigasi ke halaman masing-masing seksi:')

tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl2, ['No', 'Sub-Menu', 'Keterangan'], col_widths=[1.0, 4.0, 9.0])
data2 = [
    ('1', 'Subbagian Tata Usaha', 'Profil dan fungsi subbagian tata usaha beserta data pegawai.'),
    ('2', 'Seksi Perindustrian', 'Informasi seksi yang membidangi perindustrian.'),
    ('3', 'Seksi Perdagangan', 'Informasi seksi yang membidangi perdagangan.'),
    ('4', 'Seksi Koperasi & UKM', 'Informasi seksi yang membidangi koperasi dan usaha kecil menengah.'),
]
for i, row_data in enumerate(data2, 1):
    add_data_row(tbl2, i, row_data, center_cols=[0])
caption('Tabel 3.2 Sub-Menu Profil Organisasi')

heading('3.3.3 Menu Navigasi Admin', level=2)
body('Menu berikut hanya muncul setelah pengguna berhasil login sebagai Admin:')

tbl3 = doc.add_table(rows=6, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl3, ['No', 'Menu', 'Keterangan'], col_widths=[1.0, 3.5, 9.5])
data3 = [
    ('1', 'Dashboard', 'Halaman ringkasan berisi daftar berita terpublikasi dan statistik.'),
    ('2', 'Tambah Berita', 'Formulir penambahan berita baru yang langsung dipublikasikan.'),
    ('3', 'Review Berita', 'Antrian berita dari staff yang menunggu persetujuan atau penolakan.'),
    ('4', 'Kelola Staff', 'Manajemen akun staff: tambah, aktifkan, nonaktifkan, reset password.'),
    ('5', 'Ganti Kata Sandi / Keluar', 'Menu akun di dropdown kanan atas untuk ubah password atau logout.'),
]
for i, row_data in enumerate(data3, 1):
    add_data_row(tbl3, i, row_data, center_cols=[0])
caption('Tabel 3.3 Menu Navigasi Admin')

heading('3.3.4 Menu Navigasi Staff', level=2)
body('Menu berikut hanya muncul setelah pengguna berhasil login sebagai Staff:')

tbl4 = doc.add_table(rows=4, cols=3)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl4, ['No', 'Menu', 'Keterangan'], col_widths=[1.0, 3.5, 9.5])
data4 = [
    ('1', 'Dashboard Staff', 'Halaman ringkasan berita milik staff beserta statusnya.'),
    ('2', 'Buat Berita Baru', 'Formulir pembuatan berita baru yang akan masuk antrian review admin.'),
    ('3', 'Ganti Kata Sandi / Keluar', 'Menu akun di dropdown kanan atas untuk ubah password atau logout.'),
]
for i, row_data in enumerate(data4, 1):
    add_data_row(tbl4, i, row_data, center_cols=[0])
caption('Tabel 3.4 Menu Navigasi Staff')

# ──────────────────────────────────────────────────────────
# 3.4 Use Case Diagram
# ──────────────────────────────────────────────────────────
heading('3.4 Use Case Diagram', level=1)
body(
    'Use case diagram menggambarkan interaksi antara aktor (pengguna) dengan fungsi-fungsi '
    'yang tersedia dalam sistem. Sistem website Sudin PPKUKM melibatkan tiga aktor utama: '
    'Pengunjung (Publik), Staff, dan Admin.'
)

heading('3.4.1 Daftar Aktor', level=2)
tbl_actor = doc.add_table(rows=4, cols=3)
tbl_actor.style = 'Table Grid'
tbl_actor.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl_actor, ['No', 'Aktor', 'Deskripsi'], col_widths=[1.0, 3.0, 10.0])
actors = [
    ('1', 'Pengunjung (Publik)', 'Pengguna umum yang mengakses informasi tanpa login.'),
    ('2', 'Staff', 'Pegawai internal yang membuat dan mengelola draft berita.'),
    ('3', 'Admin', 'Super user yang mereview berita, mengelola staff, dan memiliki semua hak akses.'),
]
for i, r in enumerate(actors, 1):
    add_data_row(tbl_actor, i, r, center_cols=[0])
caption('Tabel 3.5 Daftar Aktor Sistem')

heading('3.4.2 Daftar Use Case', level=2)
tbl_uc = doc.add_table(rows=14, cols=4)
tbl_uc.style = 'Table Grid'
tbl_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl_uc, ['Kode', 'Nama Use Case', 'Aktor', 'Keterangan'], col_widths=[1.5, 4.0, 2.5, 6.0])
uc_data = [
    ('UC-01', 'Melihat Beranda',            'Publik',       'Pengunjung melihat berita terkini dan informasi utama.'),
    ('UC-02', 'Melihat Profil Organisasi',  'Publik',       'Pengunjung melihat struktur dan profil organisasi.'),
    ('UC-03', 'Melihat Sub-Seksi',          'Publik',       'Pengunjung melihat detail masing-masing seksi.'),
    ('UC-04', 'Melihat Detail Berita',      'Publik',       'Pengunjung membaca isi lengkap berita beserta media.'),
    ('UC-05', 'Melihat Layanan',            'Publik',       'Pengunjung melihat informasi layanan dan program.'),
    ('UC-06', 'Login',                      'Staff, Admin', 'Pengguna internal masuk dengan email dan password.'),
    ('UC-07', 'Logout',                     'Staff, Admin', 'Pengguna keluar dari sesi aktif.'),
    ('UC-08', 'Ganti Password',             'Staff, Admin', 'Pengguna mengubah kata sandi akun.'),
    ('UC-09', 'Buat Berita (Staff)',         'Staff',        'Staff membuat berita baru status pending untuk direview.'),
    ('UC-10', 'Edit Berita (Staff)',         'Staff',        'Staff mengedit berita miliknya yang belum dipublikasi.'),
    ('UC-11', 'Hapus Berita (Staff)',        'Staff',        'Staff menghapus berita miliknya sendiri.'),
    ('UC-12', 'Review Berita',              'Admin',        'Admin menyetujui atau menolak berita dari staff.'),
    ('UC-13', 'Kelola Berita (Admin)',       'Admin',        'Admin menambah, mengedit, dan menghapus berita.'),
]
for i, r in enumerate(uc_data, 1):
    add_data_row(tbl_uc, i, r, center_cols=[0, 2])
caption('Tabel 3.6 Daftar Use Case Sistem')

# Ubah baris terakhir (row 14) sebagai baris ke-14
add_data_row(tbl_uc, 13, ('UC-14', 'Kelola Staff', 'Admin', 'Admin menambah, mengaktifkan, menonaktifkan, dan reset password staff.'), center_cols=[0, 2])

heading('3.4.3 Deskripsi Use Case', level=2)

# UC-06 Login
body('a.  UC-06: Login', space_before=6)
tbl_uc6 = doc.add_table(rows=7, cols=2)
tbl_uc6.style = 'Table Grid'
tbl_uc6.alignment = WD_TABLE_ALIGNMENT.CENTER
for label, val in [
    ('Nama Use Case', 'Login'),
    ('Kode', 'UC-06'),
    ('Aktor', 'Staff, Admin'),
    ('Deskripsi', 'Pengguna memasukkan email dan password untuk mengakses sistem.'),
    ('Pra-kondisi', 'Pengguna belum login dan memiliki akun aktif.'),
    ('Alur Normal', '1. Pengguna membuka halaman /login.\n2. Memasukkan email dan password.\n3. Sistem memvalidasi kredensial.\n4. Sistem mengarahkan ke dashboard sesuai peran.'),
    ('Alur Alternatif', 'Jika email/password salah, sistem menampilkan pesan error. Jika akun nonaktif, akses ditolak.'),
]:
    row = tbl_uc6.add_row()
    row.cells[0].paragraphs[0].add_run(label).bold = True
    row.cells[1].paragraphs[0].add_run(val)
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                set_font(r, size=11)
tbl_uc6.rows[0].cells[0].merge(tbl_uc6.rows[0].cells[0])
# Delete the first empty row
tbl_uc6._tbl.remove(tbl_uc6.rows[0]._tr)
caption('Tabel 3.7 Deskripsi Use Case UC-06: Login')

# UC-09 Buat Berita Staff
body('b.  UC-09: Buat Berita (Staff)', space_before=6)
tbl_uc9 = doc.add_table(rows=7, cols=2)
tbl_uc9.style = 'Table Grid'
for label, val in [
    ('Nama Use Case', 'Buat Berita (Staff)'),
    ('Kode', 'UC-09'),
    ('Aktor', 'Staff'),
    ('Deskripsi', 'Staff membuat artikel berita baru yang dikirim ke antrian review admin.'),
    ('Pra-kondisi', 'Staff telah login dan berada di halaman dashboard staff.'),
    ('Alur Normal', '1. Staff memilih menu "Buat Berita Baru".\n2. Mengisi formulir: judul, kategori, isi, tanggal, lokasi, peserta.\n3. Mengunggah foto dan/atau video.\n4. Menekan tombol "Kirim".\n5. Sistem menyimpan berita dengan status pending.\n6. Admin menerima notifikasi.'),
    ('Alur Alternatif', 'Jika judul kosong atau file media tidak valid, sistem menampilkan pesan error dan formulir dipertahankan.'),
]:
    row = tbl_uc9.add_row()
    row.cells[0].paragraphs[0].add_run(label).bold = True
    row.cells[1].paragraphs[0].add_run(val)
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                set_font(r, size=11)
tbl_uc9._tbl.remove(tbl_uc9.rows[0]._tr)
caption('Tabel 3.8 Deskripsi Use Case UC-09: Buat Berita')

# UC-12 Review Berita
body('c.  UC-12: Review Berita', space_before=6)
tbl_uc12 = doc.add_table(rows=7, cols=2)
tbl_uc12.style = 'Table Grid'
for label, val in [
    ('Nama Use Case', 'Review Berita'),
    ('Kode', 'UC-12'),
    ('Aktor', 'Admin'),
    ('Deskripsi', 'Admin meninjau berita pending dari staff lalu memutuskan untuk menyetujui atau menolak.'),
    ('Pra-kondisi', 'Admin telah login. Terdapat setidaknya satu berita berstatus pending.'),
    ('Alur Normal (Setujui)', '1. Admin membuka menu "Review Berita".\n2. Melihat daftar berita pending.\n3. Memilih berita dan menekan "Setujui".\n4. Sistem mengubah status menjadi published.\n5. Berita tampil di halaman publik.'),
    ('Alur Alternatif (Tolak)', 'Admin memilih "Tolak", wajib mengisi alasan penolakan, sistem mengubah status menjadi rejected dan alasan tersimpan.'),
]:
    row = tbl_uc12.add_row()
    row.cells[0].paragraphs[0].add_run(label).bold = True
    row.cells[1].paragraphs[0].add_run(val)
    for c in row.cells:
        for p in c.paragraphs:
            for r in p.runs:
                set_font(r, size=11)
tbl_uc12._tbl.remove(tbl_uc12.rows[0]._tr)
caption('Tabel 3.9 Deskripsi Use Case UC-12: Review Berita')

# ──────────────────────────────────────────────────────────
# 3.5 Sequence Diagram
# ──────────────────────────────────────────────────────────
heading('3.5 Sequence Diagram', level=1)
body(
    'Sequence diagram menunjukkan urutan interaksi antar objek dalam sistem '
    'secara kronologis. Berikut adalah sequence diagram untuk proses-proses '
    'utama dalam sistem website Sudin PPKUKM Jakarta Barat.'
)

heading('3.5.1 Sequence Diagram: Proses Login', level=2)
body(
    'Diagram ini menjelaskan alur interaksi ketika pengguna (Staff atau Admin) '
    'melakukan proses autentikasi ke dalam sistem.'
)

diagram_box([
    '  Pengguna          Browser           Flask (app.py)         Database',
    '     |                  |                    |                   |',
    '     |-- GET /login --> |                    |                   |',
    '     |                  |-- Request -------> |                   |',
    '     |                  |<-- render login -- |                   |',
    '     |<-- Tampilkan --- |                    |                   |',
    '     |   Halaman Login  |                    |                   |',
    '     |                  |                    |                   |',
    '     |-- Isi form &---- |                    |                   |',
    '     |   POST /login -> |                    |                   |',
    '     |                  |-- Request -------> |                   |',
    '     |                  |               Validasi email           |',
    '     |                  |               format                   |',
    '     |                  |                    |-- Query User ---> |',
    '     |                  |                    |<-- User data ---- |',
    '     |                  |               Verifikasi               |',
    '     |                  |               password hash            |',
    '     |                  |               Cek role &               |',
    '     |                  |               is_active                |',
    '     |                  |                    |                   |',
    '     |                  |  [Sukses] login_user() dipanggil       |',
    '     |                  |<-- redirect Dashboard --               |',
    '     |<-- Tampilkan --- |                    |                   |',
    '     |   Dashboard      |                    |                   |',
    '     |                  |                    |                   |',
    '     |                  |  [Gagal] flash error message           |',
    '     |<-- Tampilkan --- |<-- render login -- |                   |',
    '     |   Pesan Error    |                    |                   |',
], title='Sequence Diagram 3.1 — Proses Login')
caption('Gambar 3.1 Sequence Diagram Proses Login')

heading('3.5.2 Sequence Diagram: Staff Membuat Berita', level=2)
body(
    'Diagram ini menjelaskan alur interaksi ketika staff membuat berita baru '
    'yang akan masuk ke antrian review admin.'
)
diagram_box([
    '  Staff             Browser           Flask (app.py)    Database    File System',
    '   |                   |                   |               |             |',
    '   |-- GET /staff/ --> |                   |               |             |',
    '   |   news/new        |-- Request ------> |               |             |',
    '   |                   |<-- render form -- |               |             |',
    '   |<-- Tampilkan ---- |                   |               |             |',
    '   |   Formulir Berita |                   |               |             |',
    '   |                   |                   |               |             |',
    '   |-- Isi form &----- |                   |               |             |',
    '   |   POST + files -> |                   |               |             |',
    '   |                   |-- Request ------> |               |             |',
    '   |                   |               Validasi judul      |             |',
    '   |                   |               Validasi file       |             |',
    '   |                   |               (magic number,      |             |',
    '   |                   |                MIME, ukuran)      |             |',
    '   |                   |                   |               |             |',
    '   |                   |    [File valid]   |-- Simpan File -----------> |',
    '   |                   |                   |<-- Nama File ------------- |',
    '   |                   |                   |               |             |',
    '   |                   |   Buat News(status=pending)       |             |',
    '   |                   |                   |-- INSERT News --> |         |',
    '   |                   |                   |-- INSERT Photo -> |         |',
    '   |                   |                   |-- INSERT Video -> |         |',
    '   |                   |                   |<-- Commit ------- |         |',
    '   |                   |<-- redirect ------                              |',
    '   |                   |   /staff/dashboard                              |',
    '   |<-- Dashboard ----- |  (flash: Menunggu review admin)                |',
], title='Sequence Diagram 3.2 — Staff Membuat Berita')
caption('Gambar 3.2 Sequence Diagram Staff Membuat Berita')

heading('3.5.3 Sequence Diagram: Admin Mereview Berita', level=2)
body(
    'Diagram ini menjelaskan proses admin meninjau berita yang diajukan oleh staff '
    'dan memutuskan untuk menyetujui atau menolak.'
)
diagram_box([
    '  Admin              Browser           Flask (app.py)         Database',
    '   |                    |                   |                    |',
    '   |-- GET /admin/ ---> |                   |                    |',
    '   |   reviews          |-- Request ------> |                    |',
    '   |                    |                   |-- Query pending --> |',
    '   |                    |                   |<-- List berita ---- |',
    '   |<-- Tampilkan ----- |<-- render --------                     |',
    '   |   Antrian Review   |   admin_reviews   |                    |',
    '   |                    |                   |                    |',
    '   |  [Pilih Setujui]   |                   |                    |',
    '   |-- POST /admin/ --> |                   |                    |',
    '   |   reviews/{id}/    |-- Request ------> |                    |',
    '   |   approve          |                   |-- Query News ----> |',
    '   |                    |               status = published       |',
    '   |                    |               reviewed_by = admin.id   |',
    '   |                    |               reviewed_at = now()      |',
    '   |                    |                   |-- UPDATE News ---> |',
    '   |                    |                   |<-- Commit -------- |',
    '   |<-- redirect ------- <-- redirect ------                     |',
    '   |   (flash: Berhasil dipublikasikan)                          |',
    '   |                    |                   |                    |',
    '   |  [Pilih Tolak]     |                   |                    |',
    '   |-- POST + reason -> |-- Request ------> |                    |',
    '   |                    |               [reason kosong?]         |',
    '   |                    |               -> flash error, redirect |',
    '   |                    |               [reason ada?]            |',
    '   |                    |               status = rejected        |',
    '   |                    |               reject_reason = reason   |',
    '   |                    |                   |-- UPDATE News ---> |',
    '   |                    |                   |<-- Commit -------- |',
    '   |<-- redirect -------<-- redirect -------                     |',
    '   |   (flash: Ditolak dan dikembalikan ke staff)                |',
], title='Sequence Diagram 3.3 — Admin Mereview Berita')
caption('Gambar 3.3 Sequence Diagram Admin Mereview Berita')

heading('3.5.4 Sequence Diagram: Admin Mengelola Staff', level=2)
body(
    'Diagram ini menjelaskan proses admin menambahkan akun staff baru ke dalam sistem.'
)
diagram_box([
    '  Admin              Browser           Flask (app.py)         Database',
    '   |                    |                   |                    |',
    '   |-- GET /admin/ ---> |                   |                    |',
    '   |   staff/new        |-- Request ------> |                    |',
    '   |                    |<-- render form -- |                    |',
    '   |<-- Tampilkan ----- |                   |                    |',
    '   |   Form Tambah Staff|                   |                    |',
    '   |                    |                   |                    |',
    '   |-- Isi & POST ----> |-- Request ------> |                    |',
    '   |   (username, email,|               Validasi format email    |',
    '   |    nip)            |               Cek duplikasi            |',
    '   |                    |               username/email/nip       |',
    '   |                    |                   |-- Query cek -----> |',
    '   |                    |                   |<-- Tidak ada ------ |',
    '   |                    |               Generate temp_password   |',
    '   |                    |               Hash password            |',
    '   |                    |               must_change_password=True|',
    '   |                    |                   |-- INSERT User ---> |',
    '   |                    |                   |<-- Commit -------- |',
    '   |<-- redirect -------<-- flash: password sementara xxxxxxxxx--|',
    '   |   (Admin mencatat & meneruskan password ke staff)           |',
], title='Sequence Diagram 3.4 — Admin Menambah Staff')
caption('Gambar 3.4 Sequence Diagram Admin Menambah Staff')

# ──────────────────────────────────────────────────────────
# 3.6 Activity Diagram
# ──────────────────────────────────────────────────────────
heading('3.6 Activity Diagram', level=1)
body(
    'Activity diagram menggambarkan alur aktivitas atau alur kerja (workflow) dari '
    'suatu proses dalam sistem. Berikut adalah activity diagram dari proses-proses '
    'utama sistem website Sudin PPKUKM Jakarta Barat.'
)

heading('3.6.1 Activity Diagram: Proses Login', level=2)
diagram_box([
    '                    [MULAI]',
    '                       |',
    '                       v',
    '           +-------------------------+',
    '           |  Buka Halaman /login    |',
    '           +-------------------------+',
    '                       |',
    '                       v',
    '           +-------------------------+',
    '           | Masukkan Email &        |',
    '           | Password                |',
    '           +-------------------------+',
    '                       |',
    '                       v',
    '           +-------------------------+',
    '           | Validasi Format Email   |',
    '           +-------------------------+',
    '                       |',
    '              [Valid?]-+',
    '             /         \\',
    '          Ya/           \\Tidak',
    '           /             v',
    '          |    +--------------------+',
    '          |    | Tampilkan Error    |',
    '          |    | Format Email       |',
    '          |    +--------------------+',
    '          |             |',
    '          v             v (kembali ke form)',
    ' +------------------+',
    ' | Cari User        |',
    ' | berdasarkan Email|',
    ' +------------------+',
    '          |',
    '          v',
    ' [User ditemukan?]',
    '    /         \\',
    ' Ya/           \\Tidak',
    '  /             v',
    ' |   +-----------------------+',
    ' |   | Flash: "Email/Password|',
    ' |   | salah"                |',
    ' |   +-----------------------+',
    ' v',
    ' [Password cocok?]',
    '    /         \\',
    ' Ya/           \\Tidak -> Flash Error -> Form',
    '  /',
    ' [Akun aktif?]',
    '    /         \\',
    ' Ya/           \\Tidak -> Flash "Akun dinonaktifkan"',
    '  /',
    ' login_user(user)',
    '          |',
    '          v',
    ' [must_change_password?]',
    '    /         \\',
    ' Ya/           \\Tidak',
    '  /             \\',
    ' v               v',
    '/change-password  Dashboard (sesuai role)',
    '          |',
    '        [SELESAI]',
], title='Activity Diagram 3.1 — Proses Login')
caption('Gambar 3.5 Activity Diagram Proses Login')

heading('3.6.2 Activity Diagram: Workflow Berita (Staff & Admin)', level=2)
body(
    'Diagram berikut menggambarkan alur lengkap dari pembuatan berita oleh staff '
    'hingga berita tersebut tampil di halaman publik atau dikembalikan ke staff.'
)
diagram_box([
    '   STAFF                              ADMIN                     SISTEM/DB',
    '     |                                  |                           |',
    '  [MULAI]                               |                           |',
    '     |                                  |                           |',
    '     v                                  |                           |',
    ' Buka Formulir                          |                           |',
    ' Berita Baru                            |                           |',
    '     |                                  |                           |',
    '     v                                  |                           |',
    ' Isi Data Berita                        |                           |',
    ' (judul, isi, kategori,                 |                           |',
    '  media, dll.)                          |                           |',
    '     |                                  |                           |',
    '     v                                  |                           |',
    ' [Data & Media Valid?]                  |                           |',
    '   /      \\                             |                           |',
    ' Ya/    Tidak -> Tampilkan Error        |                           |',
    '  /                                     |                           |',
    '  v                                     |                           |',
    ' Simpan Berita -----------------------> |                ------> INSERT News',
    ' (status=pending)                       |                 (status=pending)',
    '  |                                     |                           |',
    '  v                                     |                           |',
    ' Flash "Menunggu review"                |                           |',
    ' Dashboard Staff                        |                           |',
    '                                        |                           |',
    '                                    Admin Buka                     |',
    '                                    /admin/reviews                 |',
    '                                        |                           |',
    '                                        v                          |',
    '                               [Berita Pending?]                   |',
    '                                 /         \\                       |',
    '                               Ya/         Tidak -> Dashboard      |',
    '                                /                                   |',
    '                               v                                   |',
    '                        Tinjau Berita                              |',
    '                        (baca isi & media)                         |',
    '                               |                                   |',
    '                               v                                   |',
    '                        [Keputusan?]                               |',
    '                         /         \\                               |',
    '                    Setujui/       Tolak (isi alasan)             |',
    '                       /               \\                           |',
    '                      v                 v                          |',
    '            UPDATE status       UPDATE status            <---------+',
    '            = published         = rejected               |',
    '                      |                 |                |',
    '                      v                 v                |',
    '           Berita tampil          Flash "Ditolak"        |',
    '           di halaman publik      beserta alasan         |',
    '                      |                 |                |',
    '                   [SELESAI]        Staff dapat edit & kirim ulang',
], title='Activity Diagram 3.2 — Workflow Berita')
caption('Gambar 3.6 Activity Diagram Workflow Pembuatan dan Review Berita')

heading('3.6.3 Activity Diagram: Admin Mengelola Staff', level=2)
diagram_box([
    '               [MULAI]',
    '                  |',
    '                  v',
    '      +---------------------------+',
    '      |  Buka /admin/staff        |',
    '      +---------------------------+',
    '                  |',
    '                  v',
    '      +---------------------------+',
    '      |  Tampilkan Daftar Staff   |',
    '      +---------------------------+',
    '                  |',
    '           [Pilih Aksi?]',
    '         /     |     |     \\',
    '        /      |     |      \\',
    '  Tambah   Aktif/ Reset   Nonaktif',
    '  Staff    Kan   Pass.    Kan',
    '    |       |     |         |',
    '    v       v     v         v',
    ' Form   UPDATE  Generate  UPDATE',
    ' Tambah is_active temp_pw  is_active',
    ' Staff  = True   (random)  = False',
    '    |              |',
    ' Validasi    must_change',
    ' (nama,      _password',
    '  email,     = True',
    '  NIP)          |',
    '    |         Commit',
    '    v',
    ' Cek Duplikasi',
    ' username/email/NIP',
    '    |',
    ' [Unik?]',
    '  / \\',
    ' Ya/ Tidak -> Flash Error -> Kembali ke Form',
    '  /',
    ' Generate temp_password',
    ' Hash password',
    ' INSERT User (role=staff)',
    '  |',
    ' Flash password sementara',
    '  |',
    ' Redirect /admin/staff',
    '       |',
    '    [SELESAI]',
], title='Activity Diagram 3.3 — Admin Mengelola Staff')
caption('Gambar 3.7 Activity Diagram Admin Mengelola Staff')

# ──────────────────────────────────────────────────────────
# 3.7 Flowchart Sistem
# ──────────────────────────────────────────────────────────
heading('3.7 Flowchart Sistem', level=1)
body(
    'Flowchart menggambarkan alur kerja sistem secara keseluruhan dari sudut pandang '
    'proses yang dijalankan oleh sistem. Berbeda dengan activity diagram yang berfokus '
    'pada aktivitas pengguna, flowchart menunjukkan logika keputusan dan proses yang '
    'terjadi di dalam sistem.'
)

heading('3.7.1 Flowchart Utama Sistem', level=2)
diagram_box([
    '                    [START]',
    '                       |',
    '                       v',
    '           +---------------------------+',
    '           |  Pengguna Mengakses URL   |',
    '           +---------------------------+',
    '                       |',
    '                       v',
    '              [URL yang diakses?]',
    '          /         |          \\',
    '         /          |           \\',
    '  Halaman       /login,       /admin/* atau',
    '  Publik        /logout,       /staff/*',
    '  (/, /profil   /change-password',
    '   /berita, dll)|',
    '         |      v',
    '         |  [Sudah Login?]',
    '         |    /      \\',
    '         |  Ya/      Tidak -> Tampilkan',
    '         |  /                Halaman Login',
    '         v  v',
    '  +------------------+   +--------------------+',
    '  | Query Database   |   | Cek Role Pengguna  |',
    '  | (berita published|   +--------------------+',
    '  |  saja)           |           |',
    '  +------------------+   [Role == admin?]',
    '          |                /         \\',
    '          v              Ya/         Tidak (staff)',
    '  +------------------+  /             \\',
    '  | Render Template  | v               v',
    '  | HTML             | /admin/*        /staff/*',
    '  +------------------+ (full access)   (limited)',
    '          |',
    '          v',
    '   Kirim Response HTTP 200',
    '          |',
    '         [END]',
], title='Flowchart 3.1 — Alur Utama Sistem')
caption('Gambar 3.8 Flowchart Alur Utama Sistem')

heading('3.7.2 Flowchart Upload dan Validasi Media', level=2)
body(
    'Proses validasi media (foto dan video) merupakan salah satu mekanisme keamanan '
    'penting dalam sistem. Berikut adalah flowchart proses validasi dan penyimpanan file:'
)
diagram_box([
    '         [START: File diterima dari form]',
    '                       |',
    '                       v',
    '           [File kosong / tidak ada?]',
    '               /              \\',
    '             Ya/              Tidak',
    '              /                  \\',
    '  Skip (opsional)                v',
    '                      +---------------------+',
    '                      | Cek Ukuran File     |',
    '                      | Foto: maks 10 MB    |',
    '                      | Video: maks 200 MB  |',
    '                      +---------------------+',
    '                                |',
    '                        [Ukuran OK?]',
    '                          /      \\',
    '                        Ya/      Tidak -> Return Error',
    '                         /                "Ukuran melebihi batas"',
    '                        v',
    '             +---------------------+',
    '             | Cek Ekstensi File   |',
    '             | Foto: png,jpg,jpeg, |',
    '             |       gif,webp      |',
    '             | Video: mp4,webm,mov |',
    '             +---------------------+',
    '                        |',
    '                 [Ekstensi valid?]',
    '                   /         \\',
    '                 Ya/         Tidak -> Return Error',
    '                  /                  "Format tidak didukung"',
    '                 v',
    '       +--------------------+',
    '       | Cek MIME Type      |',
    '       | (Content-Type)     |',
    '       +--------------------+',
    '                  |',
    '            [MIME cocok?]',
    '              /        \\',
    '            Ya/        Tidak -> Return Error',
    '             /                  "MIME tidak valid"',
    '            v',
    '   +----------------------+',
    '   | Cek Magic Number     |',
    '   | (baca 32 byte awal): |',
    '   | PNG: 89504E47        |',
    '   | JPG: FFD8FF          |',
    '   | MP4/MOV: ftyp        |',
    '   | WebM: 1A45DFA3       |',
    '   +----------------------+',
    '              |',
    '        [Magic number OK?]',
    '           /       \\',
    '         Ya/       Tidak -> Return Error',
    '          /                 "File rusak/tidak valid"',
    '         v',
    '  +----------------------------+',
    '  | Generate nama unik:        |',
    '  | timestamp + SHA256 hash   |',
    '  +----------------------------+',
    '         |',
    '         v',
    '  +----------------------------+',
    '  | Simpan ke folder:          |',
    '  | /static/uploads/news/      |',
    '  | photos/ atau videos/       |',
    '  +----------------------------+',
    '         |',
    '         v',
    '  +----------------------------+',
    '  | Simpan URL ke Database     |',
    '  | (NewsPhoto / NewsVideo)    |',
    '  +----------------------------+',
    '         |',
    '       [END: Return nama file]',
], title='Flowchart 3.2 — Validasi dan Penyimpanan Media')
caption('Gambar 3.9 Flowchart Validasi dan Upload Media')

heading('3.7.3 Flowchart Ganti Password', level=2)
diagram_box([
    '         [START]',
    '            |',
    '            v',
    '  Pengguna buka /change-password',
    '            |',
    '            v',
    '  Masukkan: Password Lama,',
    '            Password Baru,',
    '            Konfirmasi Password',
    '            |',
    '            v',
    '  [Password lama benar?]',
    '       /         \\',
    '     Ya/         Tidak -> Flash Error',
    '      /                   "Password saat ini salah"',
    '     v',
    '  [Panjang password baru >= 8?]',
    '       /         \\',
    '     Ya/         Tidak -> Flash Error',
    '      /                   "Minimal 8 karakter"',
    '     v',
    '  [Password baru == Konfirmasi?]',
    '       /         \\',
    '     Ya/         Tidak -> Flash Error',
    '      /                   "Konfirmasi tidak cocok"',
    '     v',
    '  Hash password baru (PBKDF2:SHA256)',
    '  Simpan ke database',
    '  must_change_password = False',
    '            |',
    '            v',
    '  Flash "Password berhasil diperbarui"',
    '  Redirect ke Dashboard',
    '            |',
    '         [END]',
], title='Flowchart 3.3 — Ganti Password')
caption('Gambar 3.10 Flowchart Proses Ganti Password')

# ──────────────────────────────────────────────────────────
# 3.8 Struktur Database
# ──────────────────────────────────────────────────────────
heading('3.8 Struktur Basis Data', level=1)
body(
    'Sistem menggunakan empat tabel utama dalam basis data SQLite (dapat diubah ke MySQL). '
    'Berikut adalah struktur tabel yang digunakan:'
)

heading('3.8.1 Tabel User', level=2)
tbl_db1 = doc.add_table(rows=9, cols=4)
tbl_db1.style = 'Table Grid'
tbl_db1.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl_db1, ['Nama Field', 'Tipe Data', 'Constraint', 'Keterangan'], col_widths=[3.5, 2.5, 2.5, 5.5])
db1_data = [
    ('id',                   'Integer',  'PK, AI',         'Primary key, auto increment'),
    ('username',             'String(50)', 'Unique, NN',   'Nama pengguna (unik)'),
    ('email',                'String(100)', 'Unique, NN',  'Alamat email (unik)'),
    ('password_hash',        'String(255)', 'NN',          'Hash password PBKDF2:SHA256'),
    ('role',                 'String(20)', 'Default: staff', "Peran: 'admin' atau 'staff'"),
    ('nip',                  'String(30)', 'Unique, Nullable', 'Nomor Induk Pegawai'),
    ('is_active_account',    'Boolean',  'Default: True',  'Status aktif akun'),
    ('must_change_password', 'Boolean',  'Default: False', 'Wajib ganti password saat login'),
]
for i, r in enumerate(db1_data, 1):
    add_data_row(tbl_db1, i, r, center_cols=[1, 2])
caption('Tabel 3.10 Struktur Tabel User')

heading('3.8.2 Tabel News', level=2)
tbl_db2 = doc.add_table(rows=13, cols=4)
tbl_db2.style = 'Table Grid'
tbl_db2.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl_db2, ['Nama Field', 'Tipe Data', 'Constraint', 'Keterangan'], col_widths=[3.5, 2.5, 2.5, 5.5])
db2_data = [
    ('id',               'Integer',    'PK, AI',      'Primary key'),
    ('title',            'String(200)', 'NN',          'Judul berita'),
    ('content',          'Text',       'Nullable',    'Isi konten berita'),
    ('image_url',        'String(255)', 'Nullable',   'URL foto utama'),
    ('publish_date',     'DateTime',   'Default: now','Tanggal publikasi'),
    ('author_id',        'Integer',    'FK: user.id', 'Penulis berita'),
    ('status',           'String(20)', 'Default: pending', "Status: 'pending'/'published'/'rejected'"),
    ('category',         'String(50)', 'Default: event', "Kategori: 'ukm'/'pelatihan'/'promosi'/'event'"),
    ('start_date',       'DateTime',   'Nullable',    'Tanggal mulai kegiatan'),
    ('participant_count','Integer',    'Nullable',    'Jumlah peserta kegiatan'),
    ('location',         'String(255)','Nullable',    'Lokasi kegiatan'),
    ('reject_reason',    'Text',       'Nullable',    'Alasan penolakan oleh admin'),
]
for i, r in enumerate(db2_data, 1):
    add_data_row(tbl_db2, i, r, center_cols=[1, 2])
caption('Tabel 3.11 Struktur Tabel News')

heading('3.8.3 Tabel NewsPhoto & NewsVideo', level=2)
tbl_db3 = doc.add_table(rows=5, cols=4)
tbl_db3.style = 'Table Grid'
tbl_db3.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl_db3, ['Nama Field', 'Tipe Data', 'Constraint', 'Keterangan'], col_widths=[3.5, 2.5, 2.5, 5.5])
db3_data = [
    ('id',         'Integer',    'PK, AI',      'Primary key'),
    ('news_id',    'Integer',    'FK: news.id',  'Relasi ke tabel News'),
    ('file_url',   'String(255)', 'NN',          'Path/URL file media'),
    ('caption',    'String(255)', 'Nullable',   'Keterangan gambar/video'),
]
for i, r in enumerate(db3_data, 1):
    add_data_row(tbl_db3, i, r, center_cols=[1, 2])
caption('Tabel 3.12 Struktur Tabel NewsPhoto dan NewsVideo')

# ──────────────────────────────────────────────────────────
# 3.9 Implementasi Antarmuka
# ──────────────────────────────────────────────────────────
heading('3.9 Implementasi Antarmuka Pengguna', level=1)
body(
    'Antarmuka pengguna (User Interface) dirancang menggunakan Tailwind CSS dengan '
    'konsep glassmorphism dan dark mode. Berikut adalah deskripsi setiap halaman utama:'
)

tbl_ui = doc.add_table(rows=13, cols=3)
tbl_ui.style = 'Table Grid'
tbl_ui.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tbl_ui, ['No', 'Halaman', 'Deskripsi Komponen Antarmuka'], col_widths=[1.0, 3.5, 9.5])
ui_data = [
    ('1', 'Beranda (index.html)', 'Hero section dengan statistik UKM, breaking news ticker berjalan, filter kategori berita (tab), kartu berita dengan galeri slide, dan sidebar trending.'),
    ('2', 'Profil Organisasi (profil_organisasi.html)', 'Menampilkan struktur organisasi dengan animasi teks bergantian, foto pejabat, dan tautan ke halaman masing-masing seksi.'),
    ('3', 'Detail Berita (berita_detail.html)', 'Isi berita lengkap, galeri foto lightbox, pemutar video, serta informasi tanggal, lokasi, dan jumlah peserta.'),
    ('4', 'Halaman Layanan (layanan.html)', 'Banner Jakpreneur, daftar layanan yang tersedia, dan tautan ke program pelatihan eksternal.'),
    ('5', 'Login (login.html)', 'Form login responsif dengan toggle visibilitas password dan validasi sisi klien.'),
    ('6', 'Dashboard Admin (dashboard.html)', 'Profil admin, ringkasan statistik, aksi cepat (tambah berita, review berita, kelola staff), tabel daftar berita dengan aksi edit dan hapus.'),
    ('7', 'Dashboard Staff (staff_dashboard.html)', 'Ringkasan status berita (pending/published/rejected), tabel berita milik staff dengan aksi edit dan hapus.'),
    ('8', 'Form Berita Admin (add_news.html)', 'Formulir lengkap: judul, kategori, isi (textarea), tanggal kegiatan, lokasi, peserta, upload multi-foto, upload multi-video dengan preview.'),
    ('9', 'Form Berita Staff (staff_news_form.html)', 'Formulir serupa dengan form admin, namun berita tersimpan dengan status pending.'),
    ('10', 'Review Berita (admin_reviews.html)', 'Dua seksi: daftar berita pending dan daftar berita ditolak, dengan tombol Setujui dan Tolak (disertai input alasan).'),
    ('11', 'Kelola Staff (admin_staff_list.html)', 'Tabel daftar staff aktif/nonaktif, tombol aktif/nonaktifkan, reset password, dan tautan tambah staff baru.'),
    ('12', 'Ganti Password (change_password.html)', 'Form tiga kolom: password lama, password baru, konfirmasi password, dengan toggle visibilitas.'),
]
for i, r in enumerate(ui_data, 1):
    add_data_row(tbl_ui, i, r, center_cols=[0])
caption('Tabel 3.13 Deskripsi Antarmuka Halaman Sistem')

# ──────────────────────────────────────────────────────────
# 3.10 Keamanan Sistem
# ──────────────────────────────────────────────────────────
heading('3.10 Keamanan Sistem', level=1)
body(
    'Sistem dirancang dengan memperhatikan aspek keamanan aplikasi web. '
    'Berikut adalah mekanisme keamanan yang diterapkan:'
)
bullet('Autentikasi berbasis Flask-Login dengan session management yang aman.')
bullet('Password di-hash menggunakan algoritma PBKDF2:SHA256 dengan salt 32 byte sebelum disimpan ke database.')
bullet('Validasi dan sanitasi input pada semua formulir (email, judul, alasan, dll.).')
bullet('Validasi file upload berlapis: ukuran, ekstensi, MIME type, dan magic number (byte signature) untuk mencegah file spoofing.')
bullet('Nama file di-generate ulang secara acak (timestamp + SHA256) sehingga nama asli file tidak diekspos.')
bullet('Dekorator @admin_required dan @staff_required memastikan akses halaman sesuai role.')
bullet('HTTP security headers: X-Content-Type-Options, X-Frame-Options, X-XSS-Protection, HSTS, dan Content-Security-Policy.')
bullet('Berita dari staff tidak langsung tampil di publik — wajib melalui review admin terlebih dahulu.')

# ──────────────────────────────────────────────────────────
# SAVE
# ──────────────────────────────────────────────────────────
out_path = r'F:\Website Sudin PPKUKM\BAB_3_Laporan_Kerja_Praktek_Sudin_PPKUKM.docx'
doc.save(out_path)
print(f'File berhasil disimpan: {out_path}')
